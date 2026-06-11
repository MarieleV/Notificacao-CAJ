from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from google import genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import List, Optional
import io

app = FastAPI(title="Motor de Notificações CAJ")

# Permite que o front-end React acesse a API sem erros de CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], # Em produção, coloque a URL do seu React
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Modelos de dados esperados pelo front-end
class GerarTextoRequest(BaseModel):
    api_key: str
    textos_base: list[str]

class ExportarWordRequest(BaseModel):
    texto_final: str

@app.post("/api/gerar")
def gerar_notificacao(req: GerarTextoRequest):
    if not req.api_key:
        raise HTTPException(status_code=400, detail="API Key é obrigatória.")
    
    try:
        client = genai.Client(api_key=req.api_key)
        
        textos_juntos = "\n---\n".join(req.textos_base)
        
        # O Prompt agora força estritamente o layout que você exigiu
        prompt = f"""
        Você é o assistente jurídico do sistema. Sua tarefa é ler as infrações brutas abaixo e consolidá-las NO EXATO FORMATO exigido a seguir. 
        Não adicione saudações ou textos extras. Apenas preencha o template.
        
        Se houver mais de uma infração, some as informações nos campos "Descrição do fato gerador", "Dispositivo legal" e "Penalidade prevista" de forma coesa.
        Deixe os campos de Data, Protocolo, Funcionário e Equipe vazios.

        FORMATO OBRIGATÓRIO:
        Descrição do fato gerador: [Suas consolidações aqui]
        Dispositivo legal infringido: [Suas consolidações aqui]
        Data da constatação: 
        Protocolo: 
        Funcionário: 
        Equipe: 
        Penalidade prevista: [Suas consolidações aqui]

        INFRAÇÕES BRUTAS (SELECIONADAS PELO USUÁRIO):
        {textos_juntos}
        """
        
        resposta = client.models.generate_content(
            model='gemini-2.5-flash', 
            contents=prompt
        )
        return {"texto_gerado": resposta.text.strip()}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/exportar_word")
def exportar_word(req: ExportarWordRequest):
    doc = Document()
    
    # Configuração de Margens Padrão Jurídico
    for section in doc.sections:
        section.top_margin = Inches(1.18)
        section.left_margin = Inches(1.18)
        section.bottom_margin = Inches(0.78)
        section.right_margin = Inches(0.78)

    # Fonte Padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Corpo do Documento
    p_corpo = doc.add_paragraph()
    p_corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_corpo.paragraph_format.line_spacing = 1.5
    p_corpo.add_run(req.texto_final)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Notificacao.docx"}
    )

# ==============================================================================
# MODELOS PARA A CALCULADORA DE MULTAS
# ==============================================================================
class RateEntry(BaseModel):
    id: int
    startMonth: str
    endMonth: str
    value: str

class IrregularRow(BaseModel):
    id: int
    monthYear: str
    consumption: str
    chargedWater: str
    chargedService: str

class CalcRequest(BaseModel):
    serviceRates: List[RateEntry]
    m3Rates: List[RateEntry]
    rows: List[IrregularRow]
    aiNumber: str
    removalDate: str
    postRegM3: str
    postRegRef: str
    billedM3: str

# Funções auxiliares matemáticas em Python
def parse_brl(val: str) -> float:
    if not val: return 0.0
    try: return float(val.replace(".", "").replace(",", "."))
    except: return 0.0

def parse_my(val: str) -> Optional[datetime]:
    try: return datetime.strptime(val, "%m/%Y")
    except: return None

def in_range(target_str: str, start_str: str, end_str: str) -> bool:
    target = parse_my(target_str)
    start = parse_my(start_str)
    end = parse_my(end_str) if end_str else None
    
    if not target or not start: return False
    if end: return start <= target <= end
    return start <= target

def fmt_brl(val: float) -> str:
    return f"R$ {val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

@app.post("/api/calcular_multa")
def calcular_multa(req: CalcRequest):
    calc_rows = []
    
    valid_dates = []

    # 1. Processamento linha a linha
    for r in req.rows:
        target_dt = parse_my(r.monthYear)
        consumption = parse_brl(r.consumption)
        c_water = parse_brl(r.chargedWater)
        c_service = parse_brl(r.chargedService)
        t_charged = c_water + c_service
        
        s_rate_val = None
        m_rate_val = None
        
        if target_dt:
            # Busca taxa de serviço
            for s in req.serviceRates:
                if in_range(r.monthYear, s.startMonth, s.endMonth) and parse_brl(s.value) > 0:
                    s_rate_val = parse_brl(s.value)
                    break
            
            # Busca taxa de m3
            for m in req.m3Rates:
                if in_range(r.monthYear, m.startMonth, m.endMonth) and parse_brl(m.value) > 0:
                    m_rate_val = parse_brl(m.value)
                    break
                    
        correct_water = (consumption * m_rate_val) if (m_rate_val is not None and consumption > 0) else None
        correct_service = s_rate_val
        total_correct = (correct_water + correct_service) if (correct_water is not None and correct_service is not None) else None
        diff = (total_correct - t_charged) if total_correct is not None else None
        
        has_error = not target_dt or s_rate_val is None or m_rate_val is None
        
        if not has_error:
            valid_dates.append(target_dt)

        calc_rows.append({
            "id": r.id,
            "monthYear": r.monthYear,
            "hasError": has_error,
            "consumption": consumption,
            "correctWater": correct_water,
            "correctService": correct_service,
            "totalCorrect": total_correct,
            "chargedWater": c_water,
            "chargedService": c_service,
            "totalCharged": t_charged,
            "diff": diff,
            "m3Rate": m_rate_val
        })

    # 2. Totalizadores (KPIs)
    valid_rows = [cr for cr in calc_rows if not cr["hasError"] and cr["totalCorrect"] is not None]
    total_m3 = sum(cr["consumption"] for cr in valid_rows)
    grand_correct = sum(cr["totalCorrect"] for cr in valid_rows)
    grand_charged = sum(cr["totalCharged"] for cr in valid_rows)
    grand_diff = grand_correct - grand_charged

    # 3. Geração do Texto de Relatório
    valid_dates.sort()
    num_months = len(valid_rows)
    first_month = valid_dates[0].strftime("%m/%Y") if valid_dates else "—"
    last_month = valid_dates[-1].strftime("%m/%Y") if valid_dates else "—"

    ai_ref = f"AI {req.aiNumber.strip()}" if req.aiNumber.strip() else "[Nº do AI]"
    date_line = req.removalDate.strip() if req.removalDate.strip() else "[data]"
    post_m3 = req.postRegM3.strip() if req.postRegM3.strip() else "[m³]"
    post_ref = req.postRegRef.strip().upper() if req.postRegRef.strip() else "[MM/AAAA]"
    billed_vol = req.billedM3.strip() if req.billedM3.strip() else "[m³]"
    mes_str = "mês" if num_months == 1 else "meses"

    report_text = f"""Cálculo do consumo estimado de água ref. {ai_ref}.
Data da retirada da irregularidade: {date_line}.
{num_months} {mes_str}, com consumo impactado pela violação: {first_month} até {last_month}.
Maior consumo mês cheio lido após a regularização: {post_m3} m³ REF. {post_ref}.
Valor total do consumo estimado no período: {fmt_brl(grand_correct)}.
Valor pago pelo cliente no período da irregularidade: {fmt_brl(grand_charged)}.
Valor a ser lançado {fmt_brl(grand_diff)}.
Volume faturado no mês impactado pela violação: {billed_vol} m³.
Volume total recuperado: {total_m3} m³."""

    return {
        "rows": calc_rows,
        "totals": {
            "totalM3": total_m3,
            "grandCorrect": grand_correct,
            "grandCharged": grand_charged,
            "grandDiff": grand_diff,
            "validCount": num_months
        },
        "reportText": report_text
    }